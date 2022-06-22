#!/usr/bin/env python2
# -*- coding: utf-8 -*-
"""
Created by: Nathalia Cespedes & Maurel Mendieta
Edited by: Mauro Tassinari

"""

import gui.IntroductionWin as Introduction
import gui.SettingsWin as Settings
import gui.PlayWin as Play
import gui.TherapyWin as MainTherapy
import robotController.NAO_controller as controller
import gui.BodyWin as Body
import lib.SensorManager as manager
from PyQt4 import QtCore, QtGui
from time import time
import sys
import docx
from docx.shared import Pt



class NAO_CpWalker(object):


    def __init__(self, settings = {'UseSensors': False,
                                   'UseRobot'  : True,
                                   'RobotIp'   : "192.168.1.101",
                                   'RobotPort' : 9559
                                  }):

        #load settings
        self.settings = settings
        self.mode = None
        #Interface Objects
        self.settingsWin = Settings.SettingsWindow()
        self.IntroductionWin = Introduction.IntroductionWindow()
        self.PlayWin = Play.PlayWindow()
        self.therapyWin = MainTherapy.TherapyWindow()
        self.bodyWin = Body.BodyWindow()

        # Load signals

        # therapy win interface object

        self.settingsWin.connectIntroductionButton(self.go_to_Introduction)
        self.settingsWin.connectTherapyButton(self.start_sessionSettings)

        self.therapyWin.connectStartButton(self.onStartThreads)
        self.therapyWin.connectStopButton(self.shutdown)
        self.settingsWin.show()
        self.bodyWin.hide()
        self.PlayWin.hide()
        self.IntroductionWin.hide()
        self.therapyWin.hide()
        self.patient = None
        self.tiempo_sesion = None
        self.tiempo_inicial = None
        self.tiempo_final = None

    def go_to_Introduction(self):

        m = self.settingsWin.get_settings_data()
        self.settings['RobotIp'] = m['ip']

        if self.settings['UseRobot']:
            self.robotController = controller.RobotController({
                                                                 'name'       : "NAO",
                                                                 'ip'         : self.settings['RobotIp'],
                                                                 'port'       : self.settings['RobotPort'],
                                                                 'UseSpanish' : True,
                                                                 'MotivationTime': 300000000,
                                                                 'HeartRate_Lim': 120,
                                                                 'Cerv_Lim': 0,
                                                                 'Thor_Lim': 0
                                                              })

        self.IntroductionWin.show()
        self.settingsWin.hide()
        self.IntroductionWin.connectIntroductionButton(self.on_startIntroduction)
        self.bodyWin.hide()
        self.PlayWin.hide()

    def on_startIntroduction(self):

        self.bodyWin.hide()
        self.PlayWin.hide()
        self.IntroductionWin.show()
        """
        Start measuring therapy time
        """
        self.tiempo_inicial = time()
        print("tiempo inicial es: ", self.tiempo_inicial)

        self.patient = self.IntroductionWin.get_patient_data()
        self.robotController.patient = self.patient
        # self.robotController.tracking_faces()
        self.robotController.start_Introduction()
        self.robotController.launch_SoundTracker()
        self.robotController.load_selfPresentation()
        self.robotController.load_patientPresentation()

        self.IntroductionWin.connectnameButton(self.name_presentation)
        self.IntroductionWin.connectageButton(self.age_presentation)
        self.IntroductionWin.connectactButton(self.act_presentation)
        self.IntroductionWin.connectPassButton(self.PassToBody)

        # Set behaviors to right answer
        self.PlayWin.connectRightButton(self.robotController.load_Right)


    def speech_Production(self):
        volume = self.PlayWin.send_Volume()
        volume = ((float(volume))/100)*1.45
        print(volume)
        self.robotController.setVolume(volume)

        if self.PlayWin.get_speech()!="" and self.bodyWin.get_speech()=="":
            say = (self.PlayWin.get_speech())
        elif self.bodyWin.get_speech()!="" and self.PlayWin.get_speech()=="":
            say = (self.bodyWin.get_speech())
        else:
            say = None

        self.robotController.load_Speech(say)


    def name_presentation(self):
        self.robotController.load_name()

    def age_presentation(self):
        self.robotController.load_age()

    def act_presentation(self):
        self.robotController.load_activity()

    def Head_Explanation(self):
        self.robotController.load_HeadExplanation()

    def ArmL_Explanation(self):
        self.robotController.load_ArmLExplanation()

    def ArmR_Explanation(self):
        self.robotController.load_ArmRExplanation()

    def LegL_Explanation(self):
        self.robotController.load_LegLExplanation()

    def LegR_Explanation(self):
        self.robotController.load_LegRExplanation()

    def elephant_Game(self):
        self.robotController.load_ElephantBehavior()

    def mouse_Game(self):
        self.robotController.load_MouseBehavior()

    def monkey_Game(self):
        self.robotController.load_gorillaBehavior()

    def butterfly_Game(self):
        self.robotController.load_butterflyBehavior()

    def bye(self):
        """
        Final time of the therapy
        """
        self.tiempo_final = time()
        print("tiempo final es: ", self.tiempo_final)
        self.tiempo_sesion = self.tiempo_final - self.tiempo_inicial
        self.informe()
        self.robotController.load_bye()

    def informe(self):
        """
        Creates a report of the session in Word
        --> Change 'nombreSesion' to the desired file name. Ideal = "IDpatient - date"
        """
        self.nombreSesion = "Session 10"

        informe = docx.Document()

        title1 = informe.add_paragraph('')
        title1text = title1.add_run('RESULTS REPORT - SOCIAL INTEGRATION PROTOCOL')
        title1text.font.name = 'Times New Roman'
        title1text.font.bold = True
        title1text.font.size = Pt(14)

        title2 = informe.add_paragraph('')
        title2text = title2.add_run('Patient data')
        title2text.font.name = 'Times New Roman'
        title2text.font.underline = True
        title2text.font.size = Pt(12)

        age = informe.add_paragraph('')
        agetext = age.add_run('    Age: ' + str(self.patient['age']))
        agetext.font.name = 'Times New Roman'
        agetext.font.size = Pt(12)

        gender = informe.add_paragraph('')
        gendertext = gender.add_run('    Gender: ' + str(self.patient['gender']))
        gendertext.font.name = 'Times New Roman'
        gendertext.font.size = Pt(12)

        pathology = informe.add_paragraph('')
        pathologytext = pathology.add_run('    Pathology: ' + str(self.patient['pathology']))
        pathologytext.font.name = 'Times New Roman'
        pathologytext.font.size = Pt(12)

        id = informe.add_paragraph('')
        idtext = id.add_run('    ID: ' + str(self.patient['id']))
        idtext.font.name = 'Times New Roman'
        idtext.font.size = Pt(12)

        enter1 = informe.add_paragraph('')

        title3 = informe.add_paragraph('')
        title3text = title3.add_run('Body parts imitation')
        title3text.font.name = 'Times New Roman'
        title3text.font.underline = True
        title3text.font.size = Pt(12)

        left_arm = informe.add_paragraph('')
        if self.robotController.BI_activado:
            left_arm_text = left_arm.add_run(
                '    Attempts to imitate the left arm movement:     ' + str(self.robotController.cont_BrazoIz))
        else:
            left_arm_text = left_arm.add_run(
                '    Attempts to imitate the left arm movement:     Not selected')
        left_arm_text.font.name = 'Times New Roman'
        left_arm_text.font.size = Pt(12)

        right_arm = informe.add_paragraph('')
        if self.robotController.BD_activado:
            right_arm_text = right_arm.add_run(
                '    Attempts to imitate the right arm movement:   ' + str(self.robotController.cont_BrazoDr))
        else:
            right_arm_text = right_arm.add_run(
                '    Attempts to imitate the right arm movement:     Not selected')
        right_arm_text.font.name = 'Times New Roman'
        right_arm_text.font.size = Pt(12)

        left_leg = informe.add_paragraph('')
        if self.robotController.PI_activado:
            left_leg_text = left_leg.add_run(
                '    Attempts to imitate the left leg movement:      ' + str(self.robotController.cont_PiernaIz))
        else:
            left_leg_text = left_leg.add_run(
                '    Attempts to imitate the left leg movement:      Not selected')
        left_leg_text.font.name = 'Times New Roman'
        left_leg_text.font.size = Pt(12)

        right_leg = informe.add_paragraph('')
        if self.robotController.PD_activado:
            right_leg_text = right_leg.add_run(
                '    Attempts to imitate the right leg movement:    ' + str(self.robotController.cont_PiernaDr))
        else:
            right_leg_text = right_leg.add_run(
                '    Attempts to imitate the right leg movement:    Not selected')
        right_leg_text.font.name = 'Times New Roman'
        right_leg_text.font.size = Pt(12)

        head = informe.add_paragraph('')
        if self.robotController.cab_activado:
            head_text = head.add_run(
                '    Attempts to imitate the head movement:          ' + str(self.robotController.cont_cabeza))
        else:
            head_text = head.add_run(
                '    Attempts to imitate the head movement:          Not selected')
        head_text.font.name = 'Times New Roman'
        head_text.font.size = Pt(12)

        enter2 = informe.add_paragraph('')

        title4 = informe.add_paragraph('')
        title4text = title4.add_run('Animals imitation')
        title4text.font.name = 'Times New Roman'
        title4text.font.underline = True
        title4text.font.size = Pt(12)

        elephant = informe.add_paragraph('')
        if self.robotController.elef_activado:
            elephant_text = elephant.add_run(
                '    Attempts to imitate the elephant:           ' + str(self.robotController.cont_elefante))
        else:
            elephant_text = elephant.add_run(
                '    Attempts to imitate the elephant:             Not selected')
        elephant_text.font.name = 'Times New Roman'
        elephant_text.font.size = Pt(12)

        mouse = informe.add_paragraph('')
        if self.robotController.raton_activado:
            mouse_text = mouse.add_run(
                '    Attempts to imitate the mouse:              ' + str(self.robotController.cont_raton))
        else:
            mouse_text = mouse.add_run(
                '    Attempts to imitate the mouse:              Not selected')
        mouse_text.font.name = 'Times New Roman'
        mouse_text.font.size = Pt(12)

        butterfly = informe.add_paragraph('')
        if self.robotController.ave_activado:
            butterfly_text = butterfly.add_run(
                '    Attempts to imitate the bird:                  ' + str(self.robotController.cont_ave))
        else:
            butterfly_text = butterfly.add_run(
                '    Attempts to imitate the bird:                  Not selected')
        butterfly_text.font.name = 'Times New Roman'
        butterfly_text.font.size = Pt(12)

        gorilla = informe.add_paragraph('')
        if self.robotController.mono_activado:
            gorilla_text = gorilla.add_run(
                '    Attempts to imitate the gorilla:              ' + str(self.robotController.cont_mono))
        else:
            gorilla_text = gorilla.add_run(
                '    Attempts to imitate the gorilla:              Not selected')
        gorilla_text.font.name = 'Times New Roman'
        gorilla_text.font.size = Pt(12)

        enter3 = informe.add_paragraph('')

        time = informe.add_paragraph('')
        timetext = time.add_run('Session duration: ' + str(self.tiempo_sesion) + " sec.")
        timetext.font.name = 'Times New Roman'
        timetext.font.size = Pt(12)

        informe.save("C:\CSIC\NAO2\Sesiones\{}.docx".format(self.nombreSesion))
        # informe.save('"C:\CSIC\NAO2\Sesiones\{}.docx".format(self.nombreSesion)')

    def PassToBody(self):
        self.IntroductionWin.hide()
        self.PlayWin.hide()

        # Set body explanation
        self.bodyWin.show()
        self.robotController.load_Speech(
            '^startTag(animations/Stand, hello) Que actividad tan divertida, \\pau=200\\')
        self.robotController.load_Speech('^startTag(animations/Stand, hello) Vamos a jugar un juego, \\pau=100\\ '
                                         '^waitTag(animations/Stand, hello). '
                                         '^startTag(me)  Yo muevo  una parte del  cuerpo  y tu  me  imitas. '
                                         'Imagina que soy un espe jo!')
        self.robotController.load_standUp()
        self.bodyWin.connectHeadButton(self.Head_Explanation)
        self.bodyWin.connectArmLButton(self.ArmL_Explanation)
        self.bodyWin.connectArmRButton(self.ArmR_Explanation)
        self.bodyWin.connectLegLButton(self.LegL_Explanation)
        self.bodyWin.connectLegRButton(self.LegR_Explanation)

        #Speech production
        self.bodyWin.connectSpeechButton(self.speech_Production)
        #Set behaviors to right answer
        self.bodyWin.connectRightButton(self.robotController.load_Right)

        #Set behaviors to wrong answer
        self.bodyWin.connectWrongButton(self.robotController.load_Wrong)

        #Set change window
        self.bodyWin.connectPassButton(self.PassToPlay)
        self.bodyWin.connectPass1Button(self.go_to_Introduction)

    def PassToPlay(self):
        self.bodyWin.hide()
        self.IntroductionWin.hide()
        self.PlayWin.show()
        self.robotController.load_Speech(
            '^startTag(animations/Stand, hello) Ahora vamos a jugar a adivinar animales, \\pau=100\\')
        self.robotController.load_standUp()
        #Set animals game
        self.PlayWin.connectElephantButton(self.elephant_Game)
        self.PlayWin.connectMouseButton(self.mouse_Game)
        self.PlayWin.connectButterflyButton(self.butterfly_Game)
        self.PlayWin.connectMonkeyButton(self.monkey_Game)

        #Speech production
        self.PlayWin.connectSpeechButton(self.speech_Production)
        # Set behaviors to right answer
        self.PlayWin.connectRightButton(self.robotController.load_Right)
        #Set behaviors to wrong answer
        self.PlayWin.connectWrongButton(self.robotController.load_Wrong)

        self.PlayWin.connectPassButton(self.PassToBody)

        self.PlayWin.connectFinButton(self.bye)

    def start_sessionSettings(self):
        print('start_sessionSettings')
        m = self.settingsWin.get_settings_data()
        battery = "85%"

        self.settings['RobotIp'] = m['ip']

        if self.settings['UseRobot']:
            self.RobotCaptureThread = RobotCaptureThread(interface = self)
            self.robotController = controller.RobotController({
                                                                 'name'       : "NAO",
                                                                 'ip'         : self.settings['RobotIp'],
                                                                 'port'       : self.settings['RobotPort'],
                                                                 'UseSpanish' : True,
                                                                 'MotivationTime': 300000000,
                                                                 'HeartRate_Lim': 120,
                                                                 'Cerv_Lim': 0,
                                                                 'Thor_Lim': 0
                                                              })
        self.therapyWin.show()
        self.therapyWin.sensorsData(m['ip'], battery)

        self.sensor_Settings()

    def onStartThreads(self):

        self.SensorUpdateThread.start()
        self.SensorAcquisitionThread.start()

    def sensor_Settings(self):

        self.SensorUpdateThread  = SensorUpdateThread(f =self.sensor_update, sample = 1)

        self.Manager = manager.SensorManager( ecg   = {"port":'COM3', "sample":1},
                                              EMG   = {"ip": "192.168.1.51", "port": 30006})

        self.SensorAcquisitionThread = SensorAcquisitionThread(f=self.Manager.update_data, sample =1)

        if self.settings['UseSensors']:

            # set sensors
            self.Manager.set_sensors(ecg = True, emg=True)
            self.Manager.launch_Sensors()
            time.sleep(5)

    def sensor_update(self):

        if self.settings['UseSensors']:
            self.data = self.Manager.get_Data()
            print('Index')
            print(self.data)
            self.therapyWin.update_display_data(d = { 'hr' :self.data['ecg'],
                                                      'Inclination' : self.data['emg']['contractions']
                                                      }
                                                )
            self.therapyWin.onSensorUpdate.emit()

        else:
            self.data = self.Manager.get_data()
            self.therapyWin.update_display_data(d = {
                                                        'hr' :self.data['ecg'],
                                                        'Inclination' : self.data['emg']['contractions']
                                                      })
            self.therapyWin.onSensorUpdate.emit()


    def shutdown(self):
        self.SensorUpdateThread.shutdown()
        self.SensorAcquisitionThread.shutdown()
        self.robotController.shutdown()


class RobotCaptureThread(QtCore.QThread):
    def __init__(self, parent = None, sample = 10, interface = None):
        super(RobotCaptureThread,self).__init__()
        self.Ts = sample
        self.ON = True
        self.interface = interface



    def run(self):
        #self.interface.robotController.posture.goToPosture("StandZero", 1.0)
        while self.ON:
            d = self.interface.ManagerRx.get_data()
            self.interface.robotController.set_data(d)
            time.sleep(self.Ts)


    def shutdown(self):
        self.ON = False

class SensorUpdateThread(QtCore.QThread):

     def __init__(self, parent = None, f = None, sample = 1):
        super(SensorUpdateThread,self).__init__()
        self.f = f
        self.Ts = sample
        self.ON = True

     def run(self):

        if self.f:
            while self.ON:
                self.f()
                time.sleep(self.Ts)

     def shutdown(self):
        self.ON = False

class SensorAcquisitionThread(QtCore.QThread):

    def __init__(self, parent = None, f = None, sample = 1):
        super(SensorAcquisitionThread,self).__init__()
        self.f = f
        self.Ts = sample
        self.ON = True

    def run(self):
        if self.f:
            while self.ON:
                self.f()
                time.sleep(self.Ts)

    def shutdown(self):
        self.ON = False

    #def shutdown(self):
        #self.ON = False

if __name__ == "__main__":
    app = QtGui.QApplication(sys.argv)
    a = NAO_CpWalker()
    sys.exit(app.exec_())
